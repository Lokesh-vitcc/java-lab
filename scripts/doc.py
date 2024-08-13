import docx
from pygments import highlight
from pygments.lexers import JavaLexer
from pygments.formatters import HtmlFormatter
from htmldocx import HtmlToDocx
import os
import re


def highlight_java_code(code):
    # Highlight the code using Pygments
    lexer = JavaLexer()
    formatter = HtmlFormatter(style="colorful",noclasses=True)
    highlighted_code = highlight(code, lexer, formatter)
    return highlighted_code

doc = docx.Document()
parser = HtmlToDocx()

user_details={
    "name":"LOKESHWARAN P",
    "regNo":"23BAI1098",
}
file = "../JAVA/DataTypes.java"

def introPage(doc,user):
    doc.add_heading(f"NAME: {user_details['name']}\nREG No.: {user_details['regNo']}\n")
    doc.add_page_break() # new page

def execute(code):
    inputs = re.compile()
def addCodes(doc,files):
    parser=HtmlToDocx()

    for n,file in enumerate(files):
        with open(file) as f:
            code = f.read()
        doc.add_paragraph(f"Program {n+1}.)")
        html = highlight_java_code(code)
        parser.add_html_to_document(html,doc)
        doc.add_page_break()
    return doc

def getPaths(dir):
    return [os.path.join(dir,x) for x in os.listdir(dir)]
doc = docx.Document()
introPage(doc,user_details)
addCodes(doc,getPaths("../JAVA"))
doc.save("new.docx")